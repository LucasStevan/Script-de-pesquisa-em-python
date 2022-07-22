from numpy import var
import wikipedia
import re
from docx import document

name = input('Digite seu nome:')
wikipedia.set_lang('en')
title = input('Sobre o que quer pesquisar? \n')
while True:
    try :
        wiki= wikipedia.page(title)
        break
    except : print('Nome do projeto inválido! \n')

text = wiki.content
text = re.sub(" r '=','\n', text ")
text = re.sub(" r '=','\n', text ")
text = re.sub(" r '=','\n','\n', text")
split = text.split(" 'veja tambem' 1")
text = split[0]
print (text)

document = document()
paragraph = document.add_heading(title,0)
paragraph.alignment = 1

###Não mexe nessa bagaça não mds kkkkkk

paragraph = document.add_paragraph('  ' + text)
paragraph = document.add_paragraph(name)
paragraph.alignment = 2
document.save(title + ".docx")
input()









